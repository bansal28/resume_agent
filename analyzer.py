"""
Resume <-> Job Description analyzer.

Two layers:
  1. Pure-Python heuristics for ATS / parseability  (no API needed)
  2. An LLM pass for keyword gap analysis + bullet rewrites

The LLM provider is isolated in ONE function (`analyze_with_llm`) so you can
swap Anthropic for OpenAI / a local model without touching anything else.

Heavy deps (pdfplumber, python-docx, anthropic) are imported lazily inside the
functions that need them, so the pure helpers below stay testable on their own.
"""

import os
import re
import json

STANDARD_SECTIONS = [
    "experience", "work experience", "employment", "professional experience",
    "education", "skills", "technical skills", "projects", "summary",
    "objective", "certifications", "achievements", "publications",
]

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")

TELEGRAM_MAX = 4096


# --------------------------------------------------------------------------- #
# Parsing
# --------------------------------------------------------------------------- #
def extract_text_from_resume(path):
    """Return (text, warnings) from a .pdf or .docx resume."""
    warnings = []
    lower = path.lower()

    if lower.endswith(".pdf"):
        import pdfplumber
        parts, has_tables = [], False
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
                if page.extract_tables():
                    has_tables = True
        if has_tables:
            warnings.append(
                "Tables detected — many ATS parsers garble multi-column tables. "
                "Prefer simple, left-aligned text."
            )
        return "\n".join(parts), warnings

    if lower.endswith(".docx"):
        import docx
        d = docx.Document(path)
        paras = [p.text for p in d.paragraphs]
        if d.tables:
            warnings.append(
                "Tables detected in the .docx — ATS parsers often misread tables. "
                "Consider plain paragraphs instead."
            )
            for t in d.tables:
                for row in t.rows:
                    paras.append(" ".join(c.text for c in row.cells))
        return "\n".join(paras), warnings

    if lower.endswith(".txt"):
        with open(path, "r", encoding="utf-8", errors="ignore") as fh:
            return fh.read(), warnings

    raise ValueError("Unsupported file type. Please send a PDF, DOCX or TXT.")


# --------------------------------------------------------------------------- #
# ATS / parseability heuristics (deterministic, no API)
# --------------------------------------------------------------------------- #
def ats_format_check(text, extra_warnings=None):
    """Heuristic parseability check -> {score, issues, good}."""
    issues = list(extra_warnings or [])
    good = []

    n_words = len(text.split())
    if n_words < 150:
        issues.append(
            f"Only ~{n_words} words extracted. If your resume is longer, it may be "
            "an image/scanned PDF that an ATS can't read either. Export a text-based PDF or DOCX."
        )
    else:
        good.append(f"~{n_words} words extracted cleanly (machine-readable).")

    low = text.lower()
    found = sorted({s for s in STANDARD_SECTIONS if s in low})
    if found:
        good.append("Standard sections detected: " + ", ".join(found[:5]) + ".")
    else:
        issues.append("No standard section headings (Experience, Education, Skills...) detected.")

    if EMAIL_RE.search(text):
        good.append("Email address detected.")
    else:
        issues.append("No email address found in the parsed text.")

    if PHONE_RE.search(text):
        good.append("Phone number detected.")
    else:
        issues.append("No phone number found in the parsed text.")

    nonascii = sum(1 for c in text if ord(c) > 127)
    if text and nonascii / len(text) > 0.05:
        issues.append("High share of unusual characters — special glyphs/icons can confuse parsers.")

    score = max(0, 100 - 14 * len(issues))
    return {"score": score, "issues": issues, "good": good}


# --------------------------------------------------------------------------- #
# LLM gap analysis  (swap this one function to change provider)
# --------------------------------------------------------------------------- #
ANALYSIS_SYSTEM = """You are an expert technical recruiter and resume reviewer.
Compare the candidate's RESUME against the JOB DESCRIPTION and produce a tailoring report.

Return ONLY valid JSON (no prose, no markdown fences) with EXACTLY this shape:
{
  "match_score": <integer 0-100, how well the resume aligns with the role>,
  "match_summary": "<1-2 sentence explanation of the score>",
  "present_keywords": ["important JD skills/terms ALREADY in the resume"],
  "missing_keywords": ["important JD skills/terms NOT yet surfaced in the resume"],
  "keyword_tips": ["how/where to weave the missing terms in naturally"],
  "bullet_rewrites": [
    {"original": "<an actual line from the resume>", "improved": "<rewrite using the JD's language>"}
  ],
  "honest_gaps": ["genuine gaps the candidate cannot keyword their way around"]
}

Hard rules:
- NEVER invent experience, tools, titles, or metrics the resume does not support. Only reframe real content.
- Rewrite the 3-6 most relevant or weakest bullets.
- If a 'missing' keyword reflects a skill the candidate likely lacks, put it under honest_gaps instead.
- Be specific and concrete."""


def active_provider_model():
    """Return (provider, model) actually in use, applying defaults."""
    provider = os.environ.get("LLM_PROVIDER", "anthropic").lower()
    default = "gpt-4o-mini" if provider == "openai" else "claude-sonnet-4-6"
    return provider, os.environ.get("LLM_MODEL", default)


def llm_complete(system, user, max_tokens=2000, json_mode=True):
    """Provider-agnostic completion. json_mode=True forces JSON (for analysis);
    json_mode=False returns free text (for editing LaTeX). Returns raw text."""
    provider, model = active_provider_model()
    if provider == "openai":
        from openai import OpenAI  # lazy import
        client = OpenAI()  # reads OPENAI_API_KEY
        kwargs = {
            "model": model,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        }
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        resp = client.chat.completions.create(**kwargs)
        return resp.choices[0].message.content

    import anthropic  # lazy import
    client = anthropic.Anthropic()  # reads ANTHROPIC_API_KEY
    msg = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user}],
    )
    return "".join(b.text for b in msg.content if getattr(b, "type", None) == "text")


def analyze_with_llm(resume_text, jd_text):
    """Return the parsed analysis JSON (or {'_raw': ...} on parse failure)."""
    user = f"JOB DESCRIPTION:\n{jd_text}\n\n---\n\nRESUME:\n{resume_text}"
    raw = llm_complete(ANALYSIS_SYSTEM, user)
    try:
        return json.loads(_strip_json(raw))
    except json.JSONDecodeError:
        return {"_raw": raw}


def _strip_json(raw):
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-zA-Z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
    return raw.strip()


# --------------------------------------------------------------------------- #
# Report formatting
# --------------------------------------------------------------------------- #
def build_report(llm, ats=None):
    provider, model = active_provider_model()
    L = ["RESUME x JD ANALYSIS", f"(model: {provider} / {model})", ""]

    if "_raw" in llm:
        L.append(llm["_raw"])
    else:
        L.append(f"Match score: {llm.get('match_score', '?')}/100")
        if llm.get("match_summary"):
            L.append(llm["match_summary"])
        L.append("")

        if llm.get("missing_keywords"):
            L.append("Missing keywords (add only if true for you):")
            L += [f"  - {k}" for k in llm["missing_keywords"]]
            L.append("")
        if llm.get("keyword_tips"):
            L.append("How to weave them in:")
            L += [f"  - {t}" for t in llm["keyword_tips"]]
            L.append("")
        if llm.get("bullet_rewrites"):
            L.append("Suggested bullet rewrites:")
            for r in llm["bullet_rewrites"]:
                L.append(f"  Before: {r.get('original', '')}")
                L.append(f"  After:  {r.get('improved', '')}")
                L.append("")
        if llm.get("honest_gaps"):
            L.append("Honest gaps (keywords won't fix these):")
            L += [f"  - {g}" for g in llm["honest_gaps"]]
            L.append("")

    L.append(f"ATS / parseability check  (score: {ats['score']}/100)") if ats else None
    if ats:
        L += [f"  [ok] {g}" for g in ats["good"]]
        L += [f"  [! ] {i}" for i in ats["issues"]]
    return "\n".join(L)


def split_message(text, limit=TELEGRAM_MAX):
    """Split a long string into <=limit chunks on line boundaries (Telegram cap is 4096)."""
    chunks, cur = [], ""
    for line in text.split("\n"):
        if len(cur) + len(line) + 1 > limit:
            if cur:
                chunks.append(cur)
            cur = ""
        cur += line + "\n"
    if cur.strip():
        chunks.append(cur)
    return chunks